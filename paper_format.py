"""
paper-format - 将纯文本论文草稿一键排版为期刊标准格式 Word 文档

支持功能：
  - 自动识别论文结构（标题/摘要/关键词/章节/参考文献）
  - 等宽双栏排版（引言开始至末尾，栏间距 0.85 cm）
  - 标题/摘要/关键词保持单栏
  - 跨平台路径配置（环境变量）
"""
import shutil, os, re, tempfile
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

__version__ = '1.0.0'
__all__ = ['format_paper']

# ─────────────────────────────────────────────
# 路径配置：优先读环境变量，默认值兼容 Windows/macOS/Linux
# ─────────────────────────────────────────────
_SKILL_DIR = os.path.dirname(os.path.abspath(__file__))

TEMPLATE_PATH = os.environ.get(
    'PAPER_FORMAT_TEMPLATE',
    os.path.join(_SKILL_DIR, '论文', '论文1_明杆闸阀丝杠开度识别.docx')
)
OUTPUT_DIR = os.environ.get(
    'PAPER_FORMAT_OUTPUT',
    os.path.join(os.path.expanduser('~'), 'Desktop', '论文', 'output')
)
INPUT = os.environ.get(
    'PAPER_FORMAT_INPUT',
    os.path.join(tempfile.gettempdir(), 'paper_input.txt')
)

UNIT_WORDS = {"个", "天", "月", "年", "次", "台", "人", "根", "条", "件", "张", "座", "栋"}

STYLE = {
    'title_cn':  'Heading 1',
    'section1':  'Heading 2',
    'section2':  'Title',
    'intro':     'Heading 2',
    'kw_cn':     'Normal',
    'kw_en':     'Body Text',
    'abs_cn':    '摘要',
    'abs_en':    'Body Text',
    'body':      'Normal',
    'ref_head':  'Heading 2',
    'ref':       'List Paragraph',
}

# ─────────────────────────────────────────────
# 论文结构解析
# ─────────────────────────────────────────────

def is_section_title(l, level=None):
    if not re.match(r"^\d", l): return False
    if re.match(r"^(Abstract|Keywords)", l, re.I): return False
    if re.match(r"^\d+\.\d+\s", l): return level in (None, 'h2')
    m = re.match(r"^\d+\s+([一-鿿])", l)
    if m and m.group(1) not in UNIT_WORDS: return level in (None, 'h1')
    return False

def build_blocks(raw_lines):
    lines = [l.rstrip() for l in raw_lines if l.strip()]
    n = len(lines)

    def find(pred, default=None):
        try: return next(i for i, l in enumerate(lines) if pred(l))
        except StopIteration: return default

    title_cn_idx = 0
    en_idx = find(lambda l: l.startswith('# '))
    kw_cn_idx = find(lambda l: re.match(r'^关键词[：:]', l))
    kw_en_idx = find(lambda l: re.match(r'^Keywords', l, re.I))
    abs_cn_idx = find(lambda l: re.match(r'^摘要$', l))
    abs_en_idx = find(lambda l: re.match(r'^Abstract$', l, re.I))
    ref_head_idx = find(lambda l: l == '参考文献')
    first_section_idx = find(lambda l: is_section_title(l, 'h1'))

    blocks = []
    blocks.append(('title_cn', lines[title_cn_idx]))

    if kw_cn_idx is not None:
        blocks.append(('kw_cn', re.sub(r'^关键词[：:]', '', lines[kw_cn_idx]).strip()))
    if kw_en_idx is not None:
        blocks.append(('kw_en', re.sub(r'^Keywords[：:]?', '', lines[kw_en_idx], flags=re.I).strip()))

    before_section = first_section_idx if first_section_idx is not None else (ref_head_idx if ref_head_idx is not None else n)
    abs_en_content_idx = abs_en_idx + 1 if abs_en_idx is not None else None
    if abs_en_content_idx and abs_en_content_idx < n and not lines[abs_en_content_idx].strip():
        abs_en_content_idx += 1

    if abs_cn_idx is not None and abs_cn_idx < before_section:
        ci = abs_cn_idx + 1
        if ci < n and not lines[ci].strip(): ci += 1
        if ci < n and ci < before_section: blocks.append(('abs_cn', lines[ci]))

    if abs_en_idx is not None and abs_en_content_idx is not None and abs_en_content_idx < before_section:
        blocks.append(('abs_en', lines[abs_en_content_idx]))

    after_abs = (abs_en_content_idx + 1) if abs_en_content_idx is not None else (abs_cn_idx + 2 if abs_cn_idx is not None else before_section)
    intro_inserted = False
    for i in range(after_abs, before_section):
        s = lines[i].strip()
        if not s or is_section_title(s): continue
        if not intro_inserted: blocks.append(('intro', s)); intro_inserted = True
        else: blocks.append(('body', s))

    if first_section_idx is not None:
        i = first_section_idx
        while i < n:
            s = lines[i].strip(); i += 1
            if s == '参考文献' or re.match(r'^\[\d+\]', s): break
            if is_section_title(s, 'h1'):
                blocks.append(('section1', s))
                content_lines = []
                while i < n:
                    nxt = lines[i].strip(); ni = i + 1
                    if nxt == '参考文献' or re.match(r'^\[\d+\]', nxt) or is_section_title(nxt, 'h1'): break
                    if is_section_title(nxt, 'h2'):
                        if content_lines:
                            for cl in content_lines:
                                if cl.strip(): blocks.append(('body', cl.strip()))
                            content_lines = []
                        blocks.append(('section2', nxt))
                    else: content_lines.append(nxt)
                    i = ni
                for cl in content_lines:
                    if cl.strip(): blocks.append(('body', cl.strip()))

    if ref_head_idx is not None:
        blocks.append(('ref_head', '参考文献'))
        for i in range(ref_head_idx + 1, n):
            s = lines[i].strip()
            if re.match(r'^\[\d+\]', s): blocks.append(('ref', s))
    return blocks

# ─────────────────────────────────────────────
# 样式辅助
# ─────────────────────────────────────────────

def get_style_rPr(doc, style_name):
    try:
        style_el = doc.styles[style_name]._element
        rPr = style_el.find(qn('w:rPr'))
        return deepcopy(rPr) if rPr is not None else None
    except: return None

def apply_rPr(run, rPr_src):
    if rPr_src is None: return
    run_rPr = run._r.get_or_add_rPr()
    for child in rPr_src:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if run_rPr.find(qn('w:' + tag)) is None:
            run_rPr.append(deepcopy(child))

# ─────────────────────────────────────────────
# 核心函数
# ─────────────────────────────────────────────

def cm_to_twips(cm):
    """厘米转 twips（1 cm = 567 twips，1 inch = 1440 twips）"""
    return int(cm / 2.54 * 1440)

def format_paper(input_path, output_filename):
    """
    将论文草稿文本文件排版为 Word 文档。

    Args:
        input_path:  论文草稿文本文件路径
        output_filename: 输出 .docx 文件名（不含路径）

    Returns:
        输出文件的完整路径
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        raw_lines = f.read().split('\n')

    blocks = build_blocks(raw_lines)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    OUTPUT = os.path.join(OUTPUT_DIR, output_filename)
    shutil.copy(TEMPLATE_PATH, OUTPUT)
    doc = Document(OUTPUT)

    # 清空所有原有段落（保留样式定义）
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    # 删除模板中残留的表格
    for tbl in list(doc.element.body.iter(qn('w:tbl'))):
        tbl.getparent().remove(tbl)

    # 从输入中提取英文标题（在 kw_cn 之前的第一段全英文行）
    lines_all = [l.rstrip() for l in raw_lines if l.strip()]
    kw_cn_idx = next((i for i, l in enumerate(lines_all)
                      if re.match(r'^关键词[：:]', l)), None)
    en_title_text = None
    if kw_cn_idx is not None:
        for i in range(1, kw_cn_idx):
            s = lines_all[i].strip()
            if s and re.match(r'^[A-Z]', s) and not re.search(r'[一-鿿]', s):
                en_title_text = s; break

    # 提取各 block
    title_cn_block = next((b for b in blocks if b[0] == 'title_cn'), None)
    kw_cn_block = next((b for b in blocks if b[0] == 'kw_cn'), None)
    kw_en_block = next((b for b in blocks if b[0] == 'kw_en'), None)
    abs_cn_block = next((b for b in blocks if b[0] == 'abs_cn'), None)
    abs_en_block = next((b for b in blocks if b[0] == 'abs_en'), None)
    ref_head_block = next((b for b in blocks if b[0] == 'ref_head'), None)

    # 重组输出顺序
    new_order = []
    if title_cn_block: new_order.append(title_cn_block)
    new_order.append(('_blank', ''))
    if abs_cn_block: new_order.append(abs_cn_block)
    if kw_cn_block: new_order.append(kw_cn_block)
    new_order.append(('_blank', ''))
    if en_title_text: new_order.append(('title_en', en_title_text))
    if abs_en_block: new_order.append(abs_en_block)
    if kw_en_block: new_order.append(kw_en_block)
    excluded = {title_cn_block, kw_cn_block, kw_en_block, abs_cn_block, abs_en_block}
    for b in blocks:
        if b not in excluded and b[0] not in ('ref', 'ref_head'):
            new_order.append(b)
    if ref_head_block: new_order.append(ref_head_block)
    for b in blocks:
        if b[0] == 'ref': new_order.append(b)

    # ── XML 辅助函数 ──────────────────────────

    def set_para_spacing_para(p, before_twips=None, after_twips=None,
                                line_twips=None, first_indent_twips=None):
        pPr = p._element.get_or_add_pPr()
        sp_el = pPr.find(qn('w:spacing'))
        if sp_el is None:
            sp_el = OxmlElement('w:spacing')
            pPr.append(sp_el)
        if before_twips is not None:
            sp_el.set(qn('w:before'), str(before_twips))
        if after_twips is not None:
            sp_el.set(qn('w:after'), str(after_twips))
        if line_twips is not None:
            sp_el.set(qn('w:line'), str(line_twips))
            sp_el.set(qn('w:lineRule'), 'exact')
        if first_indent_twips is not None:
            ind_el = pPr.find(qn('w:ind'))
            if ind_el is None:
                ind_el = OxmlElement('w:ind')
                pPr.append(ind_el)
            ind_el.set(qn('w:firstLine'), str(first_indent_twips))
        def set_jc(val):
            jc_el = pPr.find(qn('w:jc'))
            if jc_el is None:
                jc_el = OxmlElement('w:jc')
                pPr.append(jc_el)
            jc_el.set(qn('w:val'), val)
        return p, set_jc

    def set_run_font(run, size_pt, font_cn, font_en='Times New Roman', bold=None):
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)
        rFonts.set(qn('w:ascii'), font_en)
        rFonts.set(qn('w:hAnsi'), font_en)
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
        if bold:
            b = rPr.find(qn('w:b'))
            if b is None:
                b = OxmlElement('w:b')
                rPr.append(b)
        return run

    # ── 按 block 类型写入段落 ─────────────────

    for btype, text in new_order:
        if btype == '_blank':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            set_para_spacing_para(p, before_twips=360, after_twips=0)
            run = p.add_run('')
            run.font.size = Pt(1)
            continue

        if btype == 'title_cn':
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 1']
            _, set_jc = set_para_spacing_para(p, before_twips=600, after_twips=210, line_twips=360)
            set_jc('center')
            run = p.add_run(text)
            set_run_font(run, 15, '黑体', bold=True)

        elif btype == 'abs_cn':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            set_para_spacing_para(p, before_twips=0, after_twips=0, line_twips=360, first_indent_twips=0)
            r1 = p.add_run('摘要：')
            set_run_font(r1, 10.5, '黑体', bold=True)
            r2 = p.add_run(text)
            set_run_font(r2, 10.5, '宋体')

        elif btype == 'kw_cn':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            set_para_spacing_para(p, before_twips=210, after_twips=0, line_twips=360)
            r1 = p.add_run('关键词：')
            set_run_font(r1, 10.5, '黑体', bold=True)
            r2 = p.add_run(text)
            set_run_font(r2, 10.5, '宋体')

        elif btype == 'title_en':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            _, set_jc = set_para_spacing_para(p, before_twips=0, after_twips=120, line_twips=360)
            set_jc('center')
            run = p.add_run(text)
            set_run_font(run, 14, '黑体', 'Times New Roman', bold=True)

        elif btype == 'abs_en':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            set_para_spacing_para(p, before_twips=0, after_twips=0, line_twips=360, first_indent_twips=0)
            r1 = p.add_run('Abstract: ')
            set_run_font(r1, 10.5, '宋体', 'Times New Roman', bold=True)
            r2 = p.add_run(text)
            set_run_font(r2, 10.5, '宋体', 'Times New Roman')

        elif btype == 'kw_en':
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            set_para_spacing_para(p, before_twips=0, after_twips=0, line_twips=360, first_indent_twips=0)
            r1 = p.add_run('Keywords: ')
            set_run_font(r1, 10.5, '宋体', 'Times New Roman', bold=True)
            r2 = p.add_run(text)
            set_run_font(r2, 10.5, '宋体', 'Times New Roman')

        elif btype == 'ref':
            p = doc.add_paragraph()
            p.style = doc.styles['List Paragraph']
            set_para_spacing_para(p, first_indent_twips=0)
            if text:
                run = p.add_run(text)
                set_run_font(run, 10.5, '宋体', 'Times New Roman')

        elif btype == 'ref_head':
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 2']
            set_para_spacing_para(p, before_twips=210, after_twips=210)
            run = p.add_run(text)
            set_run_font(run, 10.5, '宋体', 'Times New Roman')

        elif btype in ('section1', 'section2', 'intro', 'body'):
            st = STYLE.get(btype, 'Normal')
            p = doc.add_paragraph()
            p.style = doc.styles[st]
            if text:
                run = p.add_run(text)
                apply_rPr(run, get_style_rPr(doc, st))

        else:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            if text:
                run = p.add_run(text)
                set_run_font(run, 10.5, '宋体', 'Times New Roman')

    # 清理残留空段落
    for para in list(doc.paragraphs):
        if all(r.text == '' for r in para.runs):
            para._element.getparent().remove(para._element)

    # 页面设置
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.15)
        sec.bottom_margin = Cm(2.15)
        sec.left_margin = Cm(1.95)
        sec.right_margin = Cm(1.95)
        for hdr in sec.header.paragraphs:
            for r in hdr.runs: r.text = ''
        for ftr in sec.footer.paragraphs:
            for r in ftr.runs: r.text = ''

    # ── 分栏设置：正文（引言开始）双栏 ──
    n_paras = len(doc.paragraphs)
    n_blocks = len(new_order)
    offset = n_paras - n_blocks

    intro_idx = next((i for i, b in enumerate(new_order) if b[0] == 'intro'), None)

    if intro_idx is not None and intro_idx >= 1:
        pre_intro_idx = intro_idx - 1
        target_para_idx = pre_intro_idx + offset
        if 0 <= target_para_idx < n_paras:
            target_para = doc.paragraphs[target_para_idx]
            pPr = target_para._element.get_or_add_pPr()
            sectPr = OxmlElement('w:sectPr')
            pgMar = OxmlElement('w:pgMar')
            pgMar.set(qn('w:top'),    str(cm_to_twips(2.15)))
            pgMar.set(qn('w:bottom'), str(cm_to_twips(2.15)))
            pgMar.set(qn('w:left'),   str(cm_to_twips(1.95)))
            pgMar.set(qn('w:right'),  str(cm_to_twips(1.95)))
            pgMar.set(qn('w:header'), str(cm_to_twips(0.85)))
            pgMar.set(qn('w:footer'), str(cm_to_twips(0.75)))
            pgMar.set(qn('w:gutter'), '0')
            sectPr.append(pgMar)
            pPr.append(sectPr)

        last_para = doc.paragraphs[-1]
        last_pPr = last_para._element.get_or_add_pPr()
        last_sectPr = OxmlElement('w:sectPr')
        last_pgMar = OxmlElement('w:pgMar')
        last_pgMar.set(qn('w:top'),    str(cm_to_twips(2.15)))
        last_pgMar.set(qn('w:bottom'), str(cm_to_twips(2.15)))
        last_pgMar.set(qn('w:left'),   str(cm_to_twips(1.95)))
        last_pgMar.set(qn('w:right'),  str(cm_to_twips(1.95)))
        last_pgMar.set(qn('w:header'), str(cm_to_twips(0.85)))
        last_pgMar.set(qn('w:footer'), str(cm_to_twips(0.75)))
        last_pgMar.set(qn('w:gutter'), '0')
        last_sectPr.append(last_pgMar)
        last_cols = OxmlElement('w:cols')
        last_cols.set(qn('w:num'), '2')
        last_cols.set(qn('w:space'), '480')
        last_cols.set(qn('w:equalWidth'), '1')
        last_sectPr.append(last_cols)
        last_pPr.append(last_sectPr)

    doc.save(OUTPUT)
    return OUTPUT
